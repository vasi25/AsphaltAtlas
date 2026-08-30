from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

doc = Document()

# --- Page setup: A4 ---
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# --- Default style: Times New Roman 12, single space ---
style = doc.styles['Normal']
font  = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after  = Pt(0)
pf.line_spacing = Pt(12)   # single space = line height == font size

def set_font(run, bold=False, size=12):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold

def add_heading(text, level=1, size=14, bold=True, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(size)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, bold=bold, size=size)
    return p

def add_body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run)
    return p

def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(4):
    add_blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = Pt(16)
r = p.add_run('Etica Inteligenței Artificiale')
set_font(r, bold=True, size=16)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(6)
p2.paragraph_format.line_spacing = Pt(14)
r2 = p2.add_run('Provocări și Responsabilități în Era Digitală')
set_font(r2, bold=True, size=14)

add_blank()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.line_spacing = Pt(12)
r3 = p3.add_run('Referat — Probleme etice în informatică')
set_font(r3, size=12)

doc.add_page_break()

# ============================================================
# CUPRINS
# ============================================================
add_heading('Cuprins', size=14, center=False)
cuprins = [
    ('1.', 'Introducere'),
    ('2.', 'Scurt context: ascensiunea IA și impactul social'),
    ('3.', 'Principii etice fundamentale în inteligența artificială'),
    ('4.', 'Bias și discriminare algoritmică'),
    ('5.', 'Impactul inteligenței artificiale asupra pieței muncii'),
    ('6.', 'Inteligența artificială generativă — noi dileme etice'),
    ('7.', 'Cadrul de reglementare — Europa și perspectiva globală'),
    ('8.', 'Concluzii'),
    ('9.', 'Bibliografie'),
]
for nr, titlu in cuprins:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f'{nr}  {titlu}')
    set_font(r)

doc.add_page_break()

# ============================================================
# CAPITOL 1
# ============================================================
add_heading('1. Introducere', size=13)

add_body(
    'Inteligența artificială nu mai este un concept rezervat literaturii științifico-fantastice sau '
    'laboratoarelor de cercetare. În mai puțin de un deceniu, ea a pătruns în viața cotidiană cu o '
    'viteză fără precedent: recomandă ce filme să vizionăm, decide dacă suntem eligibili pentru un '
    'credit bancar, filtrează CV-urile candidaților la angajare, generează imagini, texte și coduri '
    'sursă la cerere și conduce autoturisme pe drumurile publice. Această omniprezență nu este în '
    'sine problematică — tehnologia a schimbat întotdeauna modul în care trăim și muncim.'
)
add_body(
    'Ceea ce face IA diferită de orice altă inovație tehnologică anterioară este faptul că acționează, '
    'în mod aparent, ca un agent care ia decizii. Iar acolo unde există decizii, există și consecințe '
    '— și, prin urmare, responsabilitate. Întrebările etice ridicate de inteligența artificială nu sunt '
    'abstracte sau îndepărtate. Ele se referă la cine beneficiază și cine este dezavantajat de '
    'algoritmii pe care îi folosim, la cine răspunde atunci când un sistem automat ia o decizie greșită, '
    'la ce se întâmplă cu datele noastre personale, cu locurile noastre de muncă, cu capacitatea '
    'noastră de a distinge adevărul de fals.'
)
add_body(
    'Prezentul referat își propune să exploreze principalele dimensiuni etice ale inteligenței '
    'artificiale: de la principiile teoretice care ar trebui să ghideze dezvoltarea sa, la problemele '
    'concrete de discriminare algoritmică și dezinformare, până la eforturile de reglementare la nivel '
    'european și global. Scopul nu este să condamne tehnologia, ci să înțeleagă condițiile în care '
    'aceasta poate fi dezvoltată și utilizată în mod responsabil.'
)

# ============================================================
# CAPITOL 2
# ============================================================
add_heading('2. Scurt context: ascensiunea IA și impactul social', size=13)

add_body(
    'Inteligența artificială ca domeniu academic a fost fondată în 1956, la conferința de la Dartmouth, '
    'când John McCarthy, Marvin Minsky și alți pionieri au formulat pentru prima dată ideea că procesele '
    'cognitive umane pot fi simulate de mașini. Deceniile care au urmat au alternat perioade de entuziasm '
    'cu așa-numitele „ierni ale IA" — momente în care progresul a stagnat din cauza limitărilor '
    'computaționale și a lipsei de date. Saltul decisiv a venit odată cu revoluția deep learning din '
    'jurul anului 2012, când rețelele neuronale profunde, alimentate de cantități masive de date și '
    'putere de calcul, au depășit toate abordările anterioare în domenii precum recunoașterea imaginilor '
    'și procesarea limbajului natural.'
)
add_body(
    'Dacă prima revoluție industrială a înlocuit munca fizică, iar calculatoarele clasice au automatizat '
    'sarcini repetitive și bine definite, inteligența artificială modernă este prima tehnologie care '
    'automatizează sarcini cognitive — judecăți, evaluări, predicții. Această schimbare de natură, nu '
    'doar de grad, este cea care amplifică miza etică. Un utilaj industrial poate să rănească un muncitor '
    'printr-un defect mecanic; un algoritm de angajare poate discrimina sistematic sute de mii de '
    'candidați fără ca nimeni să observe sau să răspundă.'
)
add_body(
    'Pandemia COVID-19 a accelerat dramatic adoptarea IA, forțând digitalizarea în educație, medicină, '
    'administrație și comerț. Lansarea ChatGPT în noiembrie 2022 a marcat un punct de inflexiune cultural: '
    'pentru prima dată, un sistem de IA generativă a devenit accesibil publicului larg, generând '
    'dezbateri publice intense despre impactul său asupra muncii, creativității, educației și '
    'adevărului. Această viteză a adoptării a creat un decalaj semnificativ față de capacitatea '
    'societății de a înțelege, reglementa și gestiona responsabil tehnologia.'
)

# ============================================================
# CAPITOL 3
# ============================================================
add_heading('3. Principii etice fundamentale în inteligența artificială', size=13)

add_body(
    'În ultimii ani, o serie de organizații internaționale, guverne și companii tehnologice au elaborat '
    'ghiduri și framework-uri etice pentru IA. Deși formulările diferă, un set de principii comune '
    'se regăsește în aproape toate aceste documente.'
)

add_heading('3.1. Transparență și explicabilitate', size=12, bold=True)
add_body(
    'Un sistem de IA este transparent dacă utilizatorii și cei afectați de deciziile sale pot înțelege '
    'cum funcționează și de ce produce un anumit rezultat. Aceasta este deosebit de importantă în '
    'domenii cu impact ridicat, precum medicina, justiția sau creditarea. „Problema cutiei negre" '
    '(black box problem) apare în cazul modelelor complexe de deep learning, unde nici măcar '
    'inginerii care le-au construit nu pot explica în termeni umani de ce modelul a luat o anumită '
    'decizie. Dreptul la explicație este recunoscut explicit în Regulamentul General privind Protecția '
    'Datelor (GDPR) al Uniunii Europene, care prevede că persoanele au dreptul de a nu face obiectul '
    'unor decizii bazate exclusiv pe prelucrare automată.'
)

add_heading('3.2. Corectitudine și nediscriminare', size=12, bold=True)
add_body(
    'Un sistem de IA este echitabil dacă nu produce rezultate discriminatorii pe baza unor atribute '
    'protejate precum rasa, sexul, religia sau orientarea sexuală. Corectitudinea nu este un concept '
    'simplu — există cel puțin 21 de definiții matematice ale echității, unele mutual incompatibile. '
    'Aceasta înseamnă că nu există o soluție tehnică universală: alegerea definiției de echitate este '
    'în sine o decizie valorică și, prin urmare, etică.'
)

add_heading('3.3. Responsabilitate și imputabilitate', size=12, bold=True)
p_acc = doc.add_paragraph()
p_acc.paragraph_format.space_before = Pt(0)
p_acc.paragraph_format.space_after  = Pt(6)
p_acc.paragraph_format.line_spacing = Pt(12)
p_acc.alignment = WD_ALIGN_PARAGRAPH.LEFT
r_acc = p_acc.add_run(
    'Când un sistem autonom cauzează un prejudiciu, cineva trebuie să răspundă. Responsabilitatea '
    '(accountability) în IA este complicată de faptul că între o decizie algoritmică și consecințele '
    'sale se interpun multiple straturi: cercetătorii care au dezvoltat modelul, compania care l-a '
    'antrenat, organizația care l-a implementat, utilizatorul care l-a folosit. Dispersia '
    'responsabilității poate duce la o „responsabilitate a nimănui".'
)
set_font(r_acc)

add_heading('3.4. Respectarea vieții private', size=12, bold=True)
add_body(
    'Sistemele de IA sunt consumatoare masive de date personale. Antrenarea unui model performant '
    'necesită adesea cantități uriașe de informații despre comportamentul uman — date care sunt '
    'colectate, stocate și procesate în moduri pe care utilizatorii nu le cunosc și nu le-au '
    'consimțit în mod informat. Recunoașterea facială în spații publice, analiza comportamentului '
    'online sau profilarea psihografică ridică întrebări fundamentale despre granița dintre '
    'eficiența tehnologică și dreptul la intimitate.'
)

# ============================================================
# CAPITOL 4
# ============================================================
add_heading('4. Bias și discriminare algoritmică', size=13)

add_body(
    'Una dintre cele mai documentate probleme etice ale inteligenței artificiale este tendința '
    'sistemelor de a reproduce și amplifica prejudecățile existente în societate. Bias-ul algoritmic '
    'nu este, în general, rezultatul unei intenții malițioase — el apare din datele de antrenament, '
    'din alegerile de design și din modul în care performanța modelelor este definită și măsurată.'
)
add_body(
    'Datele de antrenament reflectă lumea așa cum a fost, nu cum ar trebui să fie. Dacă în trecut '
    'anumite grupuri au fost sistematic dezavantajate — în angajare, creditare, justiție penală — '
    'modelele antrenate pe acele date vor învăța și vor reproduce acele tipare. Amazon a descoperit '
    'în 2018 că un sistem intern de IA pentru trierea CV-urilor penaliza candidatele femei, deoarece '
    'fusese antrenat pe date istorice dintr-o industrie dominată de bărbați. Compania a abandonat '
    'sistemul, dar cazul a rămas emblematic.'
)
add_body(
    'În domeniul recunoașterii faciale, studii independente — în special studiul MIT Media Lab al '
    'cercetătoarei Joy Buolamwini — au demonstrat că ratele de eroare pentru persoanele cu piele '
    'închisă la culoare și pentru femei sunt semnificativ mai mari decât pentru bărbații albi. '
    'Sistemele studiate aveau rate de eroare de până la 34% pentru femeile cu piele închisă, față '
    'de mai puțin de 1% pentru bărbații albi. Aceste sisteme sunt folosite în practică de forțele '
    'de ordine din mai multe țări, cu consecințe reale: în SUA au existat cazuri documentate de '
    'arestări eronate bazate pe identificări false prin recunoaștere facială.'
)
add_body(
    'Bias-ul algoritmic nu este inevitabil, dar combaterea lui necesită efort conștient: diversificarea '
    'echipelor care construiesc sistemele de IA, auditarea periodică a modelelor pentru efecte '
    'discriminatorii, implicarea comunităților afectate în procesul de design și adoptarea unor '
    'standarde tehnice de echitate. Fără aceste măsuri, IA riscă să devină un mecanism de '
    'amplificare a inegalităților sociale, învelit în aparența obiectivității matematice.'
)

# ============================================================
# CAPITOL 5
# ============================================================
add_heading('5. Impactul inteligenței artificiale asupra pieței muncii', size=13)

add_body(
    'Efectul automatizării asupra locurilor de muncă este una dintre cele mai dezbătute întrebări '
    'economice ale momentului. Studiul McKinsey Global Institute din 2017 estima că până în 2030 '
    'între 400 și 800 de milioane de locuri de muncă ar putea fi automatizate la nivel global. '
    'Un raport al Forumului Economic Mondial din 2023 estimează că IA va elimina 85 de milioane '
    'de locuri de muncă până în 2025, dar va crea 97 de milioane noi — un sold pozitiv, dar cu '
    'o distribuție profund inegală.'
)
add_body(
    'Problema nu este neapărat numărul total de locuri de muncă, ci distribuția lor. Locurile de '
    'muncă eliminate tind să fie concentrate în categorii bine definite — lucrători din contabilitate, '
    'transport, call center-uri, procesarea documentelor — adesea ocupate de persoane cu educație '
    'medie, din regiuni cu opțiuni alternative limitate. Locurile de muncă create tind să fie în '
    'domenii care necesită competențe tehnice avansate sau abilități social-emoționale greu de '
    'automatizat. Tranziția nu este automată: un șofer de camion din 50 de ani nu devine programator '
    'de algoritmi fără un efort masiv de recalificare.'
)
add_body(
    'Din perspectivă etică, întrebarea nu este dacă automatizarea va continua — ea va continua — '
    'ci cine suportă costurile tranziției și cine culege beneficiile. Dacă productivitatea generată '
    'de IA se concentrează la nivelul acționarilor câtorva companii, în timp ce costurile sociale '
    'ale șomajului sunt distribuite larg în populație, rezultatul este o creștere a inegalității, '
    'nu a bunăstării. Răspunsurile posibile includ impozitarea roboticii, venit universal de bază, '
    'investiții masive în recalificare și reglementarea vitezei de adoptare în sectoarele critice. '
    'Toate acestea sunt, în primul rând, decizii politice și etice, nu tehnice.'
)

# ============================================================
# CAPITOL 6
# ============================================================
add_heading('6. Inteligența artificială generativă — noi dileme etice', size=13)

add_body(
    'Apariția modelelor de IA generativă — sisteme capabile să producă text, imagini, audio și video '
    'realist la cerere — a deschis un nou capitol al problemelor etice, cu o urgență practică '
    'imediată. Dacă IA „clasică" ridica în principal probleme de discriminare și opacitate, IA '
    'generativă adaugă dileme legate de autenticitate, paternitate și integritatea informației.'
)

add_heading('6.1. Deepfake-uri și dezinformare', size=12, bold=True)
add_body(
    'Tehnologia deepfake permite generarea de înregistrări video sau audio extrem de realiste în care '
    'persoane reale spun sau fac lucruri pe care nu le-au spus sau făcut niciodată. Costul de producție '
    'al unui deepfake convingător a scăzut de la zeci de mii de dolari și echipamente specializate, la '
    'câteva minute și instrumente gratuite online. Implicațiile sunt grave: dezinformare politică, '
    'fraude financiare, hărțuire și pornografie non-consensuală. Un studiu Sensity AI estima că '
    '96% din deepfake-urile circulante online în 2019 erau conținut pornografic non-consensual, '
    'victimele fiind aproape exclusiv femei.'
)

add_heading('6.2. Dreptul de autor și paternitatea conținutului', size=12, bold=True)
add_body(
    'Modelele de IA generativă sunt antrenate pe cantități uriașe de conținut creat de oameni — '
    'cărți, articole, imagini, cod sursă — fără acordul explicit al autorilor și fără compensație. '
    'Această practică a generat valuri de procese juridice în SUA și Europa: scriitori, artiști '
    'vizuali și companii media au acționat în judecată companiile de IA pentru utilizarea operelor '
    'lor fără licență. Pe de altă parte, conținutul produs de IA ridică întrebarea inversă: cine '
    'este autorul unui text sau al unei imagini generate de un algoritm? Legislațiile actuale de '
    'copyright nu recunosc, în general, paternitatea non-umană, creând un vid juridic.'
)

add_heading('6.3. Impactul asupra educației și gândirii critice', size=12, bold=True)
add_body(
    'IA generativă pune sub presiune sistemele de evaluare academică și, mai profund, procesul de '
    'formare a gândirii critice. Atunci când un student poate genera un eseu academic în câteva '
    'secunde, întrebarea nu este doar una de fraudă academică, ci una despre scopul educației înseși. '
    'Totodată, consumul masiv de conținut generat de algoritmi — care tinde să fie fluid și plauzibil, '
    'dar nu neapărat corect sau nuanțat — poate eroda capacitatea publicului larg de a evalua '
    'informațiile în mod independent.'
)

# ============================================================
# CAPITOL 7
# ============================================================
add_heading('7. Cadrul de reglementare — Europa și perspectiva globală', size=13)

add_body(
    'Recunoașterea riscurilor asociate inteligenței artificiale a determinat guverne și organizații '
    'internaționale să elaboreze cadre de reglementare, cu viteze și abordări diferite.'
)

add_heading('7.1. Actul European privind IA (EU AI Act)', size=12, bold=True)
add_body(
    'Uniunea Europeană a adoptat în 2024 primul cadru legal comprehensiv dedicat inteligenței '
    'artificiale din lume — Regulamentul privind Inteligența Artificială (AI Act). Acesta '
    'clasifică sistemele de IA în patru categorii de risc: risc inacceptabil (sisteme interzise, '
    'precum scoring-ul social generalizat sau manipularea subliminală), risc ridicat (sisteme cu '
    'obligații stricte de transparență, audit și supraveghere umană, utilizate în domenii precum '
    'sănătatea, justiția, angajarea sau infrastructura critică), risc limitat (obligații de '
    'transparență, de exemplu chatboți care trebuie să se identifice ca IA) și risc minim '
    '(fără reglementare specifică). AI Act este un model de referință global, dar criticii '
    'atrag atenția că implementarea sa efectivă va depinde de capacitatea autorităților naționale '
    'de a înțelege și audita sisteme tehnice complexe.'
)

add_heading('7.2. Abordări globale', size=12, bold=True)
add_body(
    'Statele Unite au adoptat o abordare mai fragmentată, bazată pe ghiduri sectoriale și Executive '
    'Order-ul din octombrie 2023 privind IA, care stabilește standarde de siguranță și trasabilitate '
    'pentru sistemele cu risc ridicat, fără însă a crea un cadru legislativ unificat. China a '
    'adoptat reglementări specifice pentru algoritmii de recomandare și IA generativă, cu accent '
    'pe controlul conținutului și securitatea națională, mai mult decât pe drepturile individuale. '
    'Această divergență de abordare între marile puteri tehnologice creează riscul unui „arbitraj '
    'reglementar" — companiile își pot muta activitățile în jurisdicțiile cu cerințe mai permisive.'
)

add_heading('7.3. Limitele reglementării', size=12, bold=True)
add_body(
    'Orice cadru de reglementare se confruntă cu o problemă structurală: viteza inovației '
    'tehnologice depășește sistematic viteza proceselor legislative. Până când o lege este '
    'elaborată, dezbătută, adoptată și implementată, tehnologia reglementată s-a transformat '
    'semnificativ. Soluțiile propuse includ reglementarea principii-based (care reglementează '
    'efectele și obiectivele, nu tehnologiile specifice), crearea de autorități tehnice specializate '
    'cu capacitate rapidă de adaptare și implicarea industriei în procese de auto-reglementare '
    'monitorizate extern. Nici una dintre aceste abordări nu este perfectă; cel mai probabil, '
    'răspunsul va fi o combinație a tuturor.'
)

# ============================================================
# CAPITOL 8
# ============================================================
add_heading('8. Concluzii', size=13)

add_body(
    'Inteligența artificială este o tehnologie profund ambivalentă. Ea oferă soluții remarcabile '
    'pentru unele dintre cele mai presante probleme ale umanității — de la diagnosticarea precoce '
    'a cancerului la modelarea schimbărilor climatice, de la traducerea în timp real la '
    'accelerarea cercetării științifice. În același timp, neregulată și dezvoltată fără grijă, '
    'poate amplifica discriminarea, eroda intimitatea, destabiliza piețele muncii și submina '
    'încrederea în informație.'
)
add_body(
    'Concluzia fundamentală a analizei de față este că problemele etice ale IA nu sunt probleme '
    'tehnice cu soluții tehnice. Ele sunt probleme sociale, politice și valorice, care necesită '
    'deliberare publică, cadre juridice adecvate și o cultură a responsabilității în rândul celor '
    'care construiesc și implementează aceste sisteme. Tehnologia nu este neutră — ea codifică '
    'valori, distribuie putere și produce consecințe. Tocmai de aceea, întrebarea „ce fel de IA '
    'vrem?" este, în ultimă instanță, întrebarea „ce fel de societate vrem?"'
)
add_body(
    'Răspunsul la această întrebare nu poate fi lăsat exclusiv în seama inginerilor sau a '
    'companiilor tehnologice. El necesită participarea cetățenilor, a legiuitorilor, a cercetătorilor '
    'din domenii umaniste și sociale, a comunităților afectate. Educația în domeniul eticii '
    'tehnologice — inclusiv prin referate precum cel de față — este un prim pas modest, dar '
    'necesar, în direcția unei societăți capabile să utilizeze inteligența artificială în '
    'beneficiul tuturor.'
)

# ============================================================
# CAPITOL 9 — BIBLIOGRAFIE
# ============================================================
add_heading('9. Bibliografie', size=13)

refs = [
    'Buolamwini, J., & Gebru, T. (2018). Gender shades: Intersectional accuracy disparities in commercial gender classification. Proceedings of Machine Learning Research, 81, 1–15.',
    'European Parliament. (2024). Regulation (EU) 2024/1689 of the European Parliament and of the Council — Artificial Intelligence Act. Official Journal of the European Union.',
    'Floridi, L., et al. (2018). AI4People — An ethical framework for a good AI society: Opportunities, risks, principles, and recommendations. Minds and Machines, 28(4), 689–707.',
    'McKinsey Global Institute. (2017). A future that works: Automation, employment, and productivity. McKinsey & Company.',
    'Mittelstadt, B. D., et al. (2016). The ethics of algorithms: Mapping the debate. Big Data & Society, 3(2).',
    'O\'Neil, C. (2016). Weapons of math destruction: How big data increases inequality and threatens democracy. Crown Publishers.',
    'Russell, S., & Norvig, P. (2020). Artificial intelligence: A modern approach (4th ed.). Pearson.',
    'Sensity AI. (2019). The state of deepfakes: Landscape, threats, and impact. Sensity Research Report.',
    'World Economic Forum. (2023). Future of jobs report 2023. World Economic Forum.',
    'Zuboff, S. (2019). The age of surveillance capitalism: The fight for a human future at the new frontier of power. PublicAffairs.',
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    r = p.add_run(f'{i}. {ref}')
    set_font(r)

# ============================================================
# SAVE
# ============================================================
out = '/Users/vasi/Desktop/Etica_Inteligentei_Artificiale.docx'
doc.save(out)
print(f'Saved: {out}')
